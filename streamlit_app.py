import streamlit as st
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from io import BytesIO
import zipfile
import time
import random
from PIL import Image
from urllib.parse import urljoin, urlparse, urlunparse
import io
import re
import html
import logging
import tempfile
import os
from concurrent.futures import ThreadPoolExecutor, as_completed

# --- 1. SET PAGE CONFIG (Must be first) ---
st.set_page_config(page_title="CUIMC Web Extractor", page_icon="🩺", layout="wide")

# --- 2. INITIALIZE SESSION STATE ---
if 'history' not in st.session_state:
    st.session_state.history = []
if 'total_converted' not in st.session_state:
    st.session_state.total_converted = 0
if 'active_file' not in st.session_state:
    st.session_state.active_file = None
if 'active_name' not in st.session_state:
    st.session_state.active_name = ""
if 'bulk_zip' not in st.session_state:
    st.session_state.bulk_zip = None

# --- 3. CUIMC THEMING ---
def apply_custom_style():
    st.markdown("""
        <style>
        .stApp { background-color: #f8f9fa; }
        h1 { color: #1C3F60 !important; font-family: 'Helvetica Neue', Arial, sans-serif; font-weight: 700; }
        .stButton > button { background-color: #1C3F60; color: white; width: 100%; border-radius: 5px; border: none; font-weight: bold; }
        .stButton > button:hover { background-color: #2a5a8a; color: white; }
        .stSidebar { background-color: #e9ecef; }
        [data-testid="stMetricValue"] { color: #1C3F60; font-weight: bold; }
        </style>
    """, unsafe_allow_html=True)


# --- HTTP session with retries ---
def setup_session(retries=3, backoff_factor=0.3, status_forcelist=(429, 500, 502, 503, 504)):
    s = requests.Session()
    retry = Retry(
        total=retries,
        read=retries,
        connect=retries,
        backoff_factor=backoff_factor,
        status_forcelist=status_forcelist,
        allowed_methods=frozenset(['GET', 'POST', 'HEAD', 'OPTIONS'])
    )
    adapter = HTTPAdapter(max_retries=retry)
    s.mount('https://', adapter)
    s.mount('http://', adapter)
    s.headers.update({'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36'})
    return s

# shared session
session = setup_session()

# small helpers for image parsing
def _parse_srcset(srcset_val):
    # returns urls sorted by width if available (largest first)
    parts = [p.strip() for p in srcset_val.split(',') if p.strip()]
    candidates = []
    for p in parts:
        segs = p.split()
        url = segs[0]
        width = None
        if len(segs) > 1 and segs[1].endswith('w'):
            try:
                width = int(segs[1][:-1])
            except Exception:
                width = None
        candidates.append((url, width or 0))
    candidates.sort(key=lambda x: x[1], reverse=True)
    return [c[0] for c in candidates]

def _extract_image_candidate(img_tag, base_url):
    # try common attributes in order, prefer srcset candidates with largest width
    attrs = ['srcset', 'data-srcset', 'data-src', 'data-original', 'data-lazy', 'src']
    for a in attrs:
        val = img_tag.get(a)
        if not val:
            continue
        if a in ('srcset', 'data-srcset'):
            candidates = _parse_srcset(val)
            if candidates:
                candidate = candidates[0]
            else:
                continue
        else:
            candidate = val.split()[0]

        if candidate.startswith('data:'):
            # skip inline base64 for now
            continue

        return urljoin(base_url, candidate)
    return None

def _normalize_url(u):
    try:
        p = urlparse(u)
        p = p._replace(fragment='')
        return urlunparse(p)
    except Exception:
        return u

def scrape_images_from_page(page_url, min_w=200, min_h=150, junk_keywords=None):
    """Return list of (filename, bytes, source_url) and list of failures (url, error)"""
    if junk_keywords is None:
        junk_keywords = ['logo', 'icon', 'social', 'facebook', 'twitter', 'instagram', 'svg', 'button', 'bg', 'footer']

    results = []
    failures = []
    seen = set()

    try:
        resp = session.get(page_url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.content, 'html.parser')

        for img in soup.find_all('img'):
            try:
                img_url = _extract_image_candidate(img, page_url)
                if not img_url:
                    continue
                norm = _normalize_url(img_url)
                if norm in seen:
                    continue
                seen.add(norm)

                lower = img_url.lower()
                if any(k in lower for k in junk_keywords):
                    continue

                r = session.get(img_url, timeout=10)
                r.raise_for_status()

                try:
                    image = Image.open(io.BytesIO(r.content))
                except Exception as ee:
                    failures.append((img_url, f"PIL open failed: {ee}"))
                    continue

                w, h = image.size
                if w < min_w or h < min_h:
                    continue

                if image.mode in ("RGBA", "P"):
                    image = image.convert("RGB")

                img_buffer = io.BytesIO()
                image.save(img_buffer, format='JPEG', quality=90)
                img_buffer.seek(0)

                # derive filename from URL path
                parsed = urlparse(img_url)
                base = os.path.basename(parsed.path)
                if base:
                    name = f"{os.path.splitext(base)[0]}_{w}x{h}.jpg"
                else:
                    name = f"extracted_{w}x{h}_{len(results)}.jpg"

                results.append((name, img_buffer.getvalue(), img_url))
            except Exception as e:
                logging.exception("image extraction error")
                failures.append((img.get('src') or 'unknown', str(e)))

    except Exception as e:
        logging.exception("page fetch failed")
        failures.append((page_url, str(e)))

    return results, failures


# --- 4. SCRAPING & FORMATTING LOGIC ---
def extract_content(url, retries=2):
    """
    Extract textual content from a page and return (title, formatted_data).
    formatted_data: list of chunks {'tag': tag, 'content': [(type, value), ...]}
    types: 'text', 'bold', 'italic', 'link'
    """
    attempt = 0
    while attempt <= retries:
        try:
            time.sleep(random.uniform(0.2, 0.5))
            response = session.get(url, timeout=(5, 12))

            if response.status_code == 429:
                if attempt < retries:
                    time.sleep(5)
                    attempt += 1
                    continue
                return None, "RATE_LIMIT_ERROR"

            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')

            page_title = soup.find('h1')
            if page_title:
                title_text = page_title.get_text().strip()
            else:
                url_parts = [p for p in url.split('/') if p]
                title_text = url_parts[-1] if url_parts else "Extracted_Page"

            for element in soup(["script", "style", "nav", "footer", "header", "form", "iframe", "noscript"]):
                element.decompose()

            content_area = soup.find('main') or soup.find('article') or soup.body
            formatted_data = []
            tags_to_save = ['p', 'h1', 'h2', 'h3', 'h4', 'li', 'blockquote', 'figure']

            for element in content_area.find_all(tags_to_save):
                if element.find_parent(tags_to_save):
                    continue

                chunk = {'tag': element.name, 'content': []}

                for child in element.children:
                    if isinstance(child, NavigableString):
                        text_content = re.sub(r'\s+', ' ', str(child)).strip()
                        if text_content:
                            chunk['content'].append(('text', text_content))
                    else:
                        name = (child.name or '').lower()
                        if name in ['b', 'strong']:
                            txt = child.get_text(separator=' ', strip=True)
                            if txt:
                                chunk['content'].append(('bold', re.sub(r'\s+', ' ', txt)))
                        elif name in ['em', 'i']:
                            txt = child.get_text(separator=' ', strip=True)
                            if txt:
                                chunk['content'].append(('italic', re.sub(r'\s+', ' ', txt)))
                        elif name == 'a':
                            link_text = child.get_text(separator=' ', strip=True)
                            href = child.get('href')
                            if link_text:
                                chunk['content'].append(('link', (re.sub(r'\s+', ' ', link_text), href)))
                        else:
                            txt = child.get_text(separator=' ', strip=True)
                            if txt:
                                chunk['content'].append(('text', re.sub(r'\s+', ' ', txt)))

                # check for non-empty content
                has_text = any(
                    (t == 'text' and str(v).strip()) or (t in ('bold', 'italic') and str(v).strip()) or (t == 'link' and v[0].strip())
                    for t, v in chunk['content']
                ) if chunk['content'] else False

                if has_text:
                    formatted_data.append(chunk)

            return title_text, formatted_data

        except Exception as e:
            logging.exception('extract_content')
            if attempt < retries:
                time.sleep(2)
                attempt += 1
            else:
                return None, str(e)

def create_word_doc(title, formatted_data):
    doc = Document()
    doc.add_heading(title, 0)

    for chunk in formatted_data:
        tag = chunk.get('tag', '')

        # Use heading styles for h1-h4
        if tag and tag.startswith('h') and len(tag) == 2 and tag[1].isdigit():
            level = min(3, int(tag[1]))
            heading_text = ' '.join((v if t != 'link' else v[0]) for t, v in chunk['content'])
            doc.add_heading(heading_text.strip(), level)
            continue

        p_style = None
        if tag == 'li':
            p_style = 'List Bullet'
        elif tag == 'blockquote':
            p_style = 'Intense Quote'

        p = doc.add_paragraph(style=p_style) if p_style else doc.add_paragraph()

        for style_type, val in chunk['content']:
            if style_type == 'link':
                link_text, href = val
                run = p.add_run(link_text)
                run.italic = True
                if href:
                    p.add_run(f" ({href})")
            else:
                run = p.add_run(val)
                if style_type == 'bold':
                    run.bold = True
                if style_type == 'italic':
                    run.italic = True

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def create_html(title, formatted_data):
    """Return a BytesIO containing a simple HTML rendering of the extracted content."""
    def _render(content):
        out = []
        for t, v in content:
            if t == 'text':
                out.append(html.escape(v))
            elif t == 'bold':
                out.append(f"<strong>{html.escape(v)}</strong>")
            elif t == 'italic':
                out.append(f"<em>{html.escape(v)}</em>")
            elif t == 'link':
                txt, href = v
                if href:
                    out.append(f"<a href=\"{html.escape(href)}\" target=\"_blank\" rel=\"noopener noreferrer\">{html.escape(txt)}</a>")
                else:
                    out.append(html.escape(txt))
            else:
                out.append(html.escape(str(v)))
        return ''.join(out)

    parts = []
    parts.append('<!doctype html>')
    parts.append('<html lang="en">')
    parts.append('<head>')
    parts.append('<meta charset="utf-8">')
    parts.append(f'<title>{html.escape(title)}</title>')
    parts.append('</head>')
    parts.append('<body>')
    parts.append(f'<h1>{html.escape(title)}</h1>')

    in_list = False
    for chunk in formatted_data:
        tag = chunk.get('tag', '')
        if tag == 'li':
            if not in_list:
                parts.append('<ul>')
                in_list = True
            parts.append(f"<li>{_render(chunk['content'])}</li>")
            continue
        else:
            if in_list:
                parts.append('</ul>')
                in_list = False

        if tag and tag.startswith('h') and len(tag) == 2 and tag[1].isdigit():
            level = min(4, int(tag[1]))
            parts.append(f"<h{level}>{_render(chunk['content'])}</h{level}>")
            continue

        if tag == 'blockquote':
            parts.append(f"<blockquote>{_render(chunk['content'])}</blockquote>")
            continue

        parts.append(f"<p>{_render(chunk['content'])}</p>")

    if in_list:
        parts.append('</ul>')

    parts.append('</body></html>')

    html_str = '\n'.join(parts)
    bio = BytesIO()
    bio.write(html_str.encode('utf-8'))
    bio.seek(0)
    return bio

def clean_filename(title):
    clean = "".join([c for c in title if c.isalnum() or c==' ']).strip().replace(' ', '_')
    return clean if clean else "extracted_content"

# --- 5. APP LAYOUT ---
apply_custom_style()

st.title("🩺 CUIMC Web Extractor")

with st.sidebar:
    st.header("📊 Dashboard")
    st.metric("Total Processed", st.session_state.total_converted)
    st.divider()
    st.header("📜 Session History")
    for item in reversed(st.session_state.history):
        st.write(f"• {item}")
    if st.button("Clear All Data"):
        st.session_state.history = []
        st.session_state.total_converted = 0
        st.session_state.active_file = None
        st.session_state.bulk_zip = None
        st.rerun()

# --- 4 TABS ---
tab1, tab2, tab3, tab4 = st.tabs(["📄 Single URL (Word)", "📦 Bulk ZIP (Word)", "🖼️ Extract Images (ZIP)", "🗂️ Extract ALL (Word + Images)"])

with tab1:
    url_input = st.text_input("Paste target URL:", key="single_input")
    if st.button("Generate Word Document", key="btn_single"):
        with st.spinner("Processing..."):
            title, data = extract_content(url_input)
            if data == "RATE_LIMIT_ERROR":
                st.error("⚠️ Server is rate-limiting us. Please wait 60 seconds.")
            elif data and isinstance(data, list):
                st.session_state.active_file = create_word_doc(title, data)
                st.session_state.active_name = f"{clean_filename(title)}.docx"
                # also prepare an HTML export
                st.session_state.active_html = create_html(title, data)
                st.session_state.active_html_name = f"{clean_filename(title)}.html"
                st.session_state.total_converted += 1
                if title not in st.session_state.history:
                    st.session_state.history.append(title)
            else:
                st.error(f"Error: {data}")

    if st.session_state.active_file:
        st.success(f"✅ Ready: {st.session_state.active_name}")
        st.download_button(
            label="📥 Download Word File",
            data=st.session_state.active_file,
            file_name=st.session_state.active_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        if 'active_html' in st.session_state and st.session_state.active_html:
            st.download_button(
                label="🌐 Download HTML File",
                data=st.session_state.active_html,
                file_name=st.session_state.active_html_name,
                mime="text/html"
            )

_NON_HTML_EXTS = {
    '.pdf', '.jpg', '.jpeg', '.png', '.gif', '.svg', '.webp', '.avif',
    '.zip', '.gz', '.tar', '.docx', '.xlsx', '.pptx', '.csv',
    '.mp4', '.mp3', '.avi', '.mov', '.wav', '.ogg',
    '.ico', '.json', '.xml', '.js', '.css',
}

def is_likely_html(url):
    """Return True if the URL extension suggests an HTML page (or no extension)."""
    try:
        path = urlparse(url).path.lower().rstrip('/')
        _, ext = os.path.splitext(path)
        return ext not in _NON_HTML_EXTS
    except Exception:
        return True

with tab2:
    bulk_input = st.text_area("Paste URLs (one per line):", height=200)
    col_conc, col_html = st.columns([2, 1])
    with col_conc:
        concurrency = st.slider("Concurrent requests", min_value=1, max_value=10, value=5,
                                help="Higher = faster, but more likely to trigger rate limits on some sites")
    with col_html:
        html_only = st.checkbox("HTML pages only", value=True,
                                help="Skip PDFs, images, and other non-HTML URLs")

    if st.button("Process Bulk List", key="btn_bulk"):
        raw_list = [u.strip() for u in bulk_input.split('\n') if u.strip()]

        # Partition into HTML and skipped
        if html_only:
            url_list = [u for u in raw_list if is_likely_html(u)]
            skipped = [u for u in raw_list if not is_likely_html(u)]
        else:
            url_list = raw_list
            skipped = []

        if skipped:
            st.info(f"Skipping {len(skipped)} non-HTML URL(s): " + ", ".join(s[:60] for s in skipped[:5]) + ("…" if len(skipped) > 5 else ""))

        if url_list:
            # ── per-URL status tracking ──────────────────────────────────
            # States: 0=pending, 1=fetching, 2=done, 3=failed
            import threading
            statuses = [0] * len(url_list)   # 0=pending 1=fetching 2=done 3=failed
            lock = threading.Lock()

            ICONS = {0: '⬜', 1: '🔄', 2: '✅', 3: '❌'}

            progress_bar = st.progress(0)
            summary_text = st.empty()
            grid_placeholder = st.empty()

            def render_grid():
                with lock:
                    snap = list(statuses)
                done   = snap.count(2)
                failed = snap.count(3)
                fetch  = snap.count(1)
                total  = len(snap)
                progress_bar.progress((done + failed) / total,
                                      text=f"{done + failed}/{total} done  •  {fetch} in-flight  •  {failed} failed")
                # show a compact icon grid (up to 60 icons before truncating)
                icons = ''.join(ICONS[s] for s in snap[:60])
                if total > 60:
                    icons += f'  +{total - 60} more'
                grid_placeholder.markdown(icons)

            def fetch_url_tracked(idx_url):
                idx, url = idx_url
                with lock:
                    statuses[idx] = 1        # fetching
                render_grid()
                title, data = extract_content(url)
                with lock:
                    statuses[idx] = 2 if (data and isinstance(data, list)) else 3
                render_grid()
                return idx, url, title, data

            render_grid()
            results_map = {}
            failed_urls = []

            zip_buffer = BytesIO()
            zipf = zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED)
            try:
                with ThreadPoolExecutor(max_workers=concurrency) as executor:
                    futures = [executor.submit(fetch_url_tracked, (i, u))
                               for i, u in enumerate(url_list)]
                    for future in as_completed(futures):
                        idx, url, title, data = future.result()
                        if data and isinstance(data, list):
                            results_map[idx] = (url, title, data)
                        else:
                            failed_urls.append({'url': url, 'error': data})

                # pack ZIP in original URL order
                success_count = 0
                for idx in sorted(results_map.keys()):
                    url, title, data = results_map[idx]
                    doc_io = create_word_doc(title, data)
                    html_io = create_html(title, data)
                    safe_base = f"{idx + 1:02d}_{clean_filename(title)}"
                    zipf.writestr(f"{safe_base}.docx", doc_io.getvalue())
                    zipf.writestr(f"{safe_base}.html", html_io.getvalue())
                    st.session_state.total_converted += 1
                    if title not in st.session_state.history:
                        st.session_state.history.append(title)
                    success_count += 1
            finally:
                zipf.close()

            grid_placeholder.empty()

            if success_count > 0:
                st.session_state.bulk_zip = zip_buffer.getvalue()
                st.success(f"✅ Successfully processed {success_count} of {len(url_list)} URLs")
                if failed_urls:
                    with st.expander(f"❌ {len(failed_urls)} failed URL(s)"):
                        for item in failed_urls:
                            st.write(f"- {item['url']}: {item['error']}")
            else:
                st.error("All URLs failed — check that they are reachable HTML pages.")
        elif not skipped:
            st.warning("No URLs to process.")

    if st.session_state.bulk_zip:
        st.download_button(
            label="📥 Download ZIP Archive",
            data=st.session_state.bulk_zip,
            file_name="cuimc_batch_files.zip",
            mime="application/zip"
        )

with tab3:
    st.header("🖼️ Extract & Convert Images")
    st.markdown("Scrape images, convert WebP/PNG to JPG, and download as a dynamically named ZIP.")

    target_url_img = st.text_input("Enter Page URL to Scrape:", key="img_input")
    
    with st.expander("⚙️ Adjust Scraping Filters", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            min_width = st.number_input("Minimum Width (px)", value=200, step=50)
        with col2:
            min_height = st.number_input("Minimum Height (px)", value=150, step=50)

    if st.button("🔍 Extract Images", type="primary", key="btn_img"):
        if target_url_img:
            with st.spinner("Scraping page, filtering junk, and packing ZIP file..."):
                # Use the shared scraper helper
                img_results, img_failures = scrape_images_from_page(target_url_img, min_w=min_width, min_h=min_height)

                # Try to name the ZIP after the page
                try:
                    sampled_title, _ = extract_content(target_url_img)
                    page_name = sampled_title or target_url_img
                except Exception:
                    url_parts = [p for p in target_url_img.split('/') if p]
                    page_name = url_parts[-1] if url_parts else "Images"

                zip_filename = f"{clean_filename(page_name)}.zip"

                if not img_results:
                    st.warning("No images found matching criteria.")
                    if img_failures:
                        with st.expander("Failures during image extraction"):
                            for u, err in img_failures:
                                st.write(f"- {u}: {err}")
                else:
                    st.success(f"✅ Extracted {len(img_results)} images. ({len(img_failures)} failures)")

                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        # add images
                        for file_name, img_bytes, src in img_results:
                            zip_file.writestr(file_name, img_bytes)

                        # add manifest
                        manifest_lines = [f"{file_name} -> {src}" for file_name, _, src in img_results]
                        if img_failures:
                            manifest_lines.append("\nFailures:")
                            manifest_lines += [f"{u} -> {err}" for u, err in img_failures]
                        zip_file.writestr('manifest.txt', "\n".join(manifest_lines))

                    st.download_button(
                        label=f"📦 Download {zip_filename}",
                        data=zip_buffer.getvalue(),
                        file_name=zip_filename,
                        mime="application/zip",
                        type="primary",
                        use_container_width=True
                    )

# ==========================================
# TAB 4: THE GOD MODE (WORD + IMAGES)
# ==========================================
with tab4:
    st.header("🗂️ Extract ALL (Word + Images)")
    st.markdown("Rip the formatted text AND all clinical images into a single, perfectly organized ZIP file.")
    
    target_url_all = st.text_input("Enter Page URL:", key="all_input")
    
    with st.expander("⚙️ Image Scraping Filters", expanded=False):
        colA, colB = st.columns(2)
        with colA:
            min_w = st.number_input("Minimum Image Width (px)", value=200, step=50, key="w_all")
        with colB:
            min_h = st.number_input("Minimum Image Height (px)", value=150, step=50, key="h_all")

    if st.button("🚀 Extract Full Page", type="primary", key="btn_all"):
        if target_url_all:
            with st.spinner("Scraping text, converting images, and building your master ZIP..."):
                try:
                    # 1. Grab the Text
                    title, data = extract_content(target_url_all)
                    if data == "RATE_LIMIT_ERROR" or not isinstance(data, list):
                        st.error("Text extraction failed or rate limited.")
                        st.stop()
                        
                    doc_io = create_word_doc(title, data)
                    safe_title = clean_filename(title)
                    master_zip_name = f"{safe_title}_Full_Export.zip"
                    
                    # 2. Grab the Images using the shared helper
                    extracted_images, image_failures = scrape_images_from_page(target_url_all, min_w=min_w, min_h=min_h)

                    # 3. Build the Master ZIP
                    zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        # Write the Word Doc
                        zip_file.writestr(f"{safe_title}.docx", doc_io.getvalue())
                        # include an HTML export as well
                        try:
                            html_io = create_html(title, data)
                            zip_file.writestr(f"{safe_title}.html", html_io.getvalue())
                        except Exception:
                            # non-fatal: continue packaging
                            pass

                        # Write the Images into an 'images' folder
                        for file_name, img_bytes, src in extracted_images:
                            zip_file.writestr(f"images/{file_name}", img_bytes)

                        # add manifest mapping
                        manifest_lines = [f"images/{file_name} -> {src}" for file_name, _, src in extracted_images]
                        if image_failures:
                            manifest_lines.append("\nFailures:")
                            manifest_lines += [f"{u} -> {err}" for u, err in image_failures]
                        zip_file.writestr('manifest.txt', "\n".join(manifest_lines))

                    st.success(f"✅ Extracted '{title}' and {len(extracted_images)} images. ({len(image_failures)} failures)")
                    st.session_state.total_converted += 1

                    st.download_button(
                        label=f"📦 Download Master ZIP ({safe_title})",
                        data=zip_buffer.getvalue(),
                        file_name=master_zip_name,
                        mime="application/zip",
                        type="primary",
                        use_container_width=True
                    )
                    
                except Exception as e:
                    st.error(f"Full extraction failed. Error: {e}")
